"""
Document Parsers
Implements both rule-based and AI-based document parsing strategies.
Supports PDF, DOCX, HTML, and plain text files.
"""

import os
import re
from abc import ABC, abstractmethod
from typing import List, Dict, Any, Optional
from pathlib import Path
from dataclasses import dataclass, field
import logging

logger = logging.getLogger(__name__)


@dataclass
class Document:
    """Represents a parsed document."""
    content: str
    metadata: Dict[str, Any] = field(default_factory=dict)
    source: str = ""
    doc_type: str = ""
    
    def __post_init__(self):
        if not self.metadata:
            self.metadata = {}
        self.metadata["source"] = self.source
        self.metadata["doc_type"] = self.doc_type


class DocumentParser(ABC):
    """Abstract base class for document parsers."""
    
    @abstractmethod
    def parse(self, file_path: str) -> Document:
        """Parse a document and return its content."""
        pass
    
    @abstractmethod
    def supports(self, file_path: str) -> bool:
        """Check if this parser supports the given file type."""
        pass


class PDFParser(DocumentParser):
    """
    PDF Document Parser
    Uses pypdf for rule-based parsing with optional AI enhancement.
    """
    
    SUPPORTED_EXTENSIONS = [".pdf"]
    
    def supports(self, file_path: str) -> bool:
        return Path(file_path).suffix.lower() in self.SUPPORTED_EXTENSIONS
    
    def parse(self, file_path: str) -> Document:
        """Parse PDF document using pypdf."""
        try:
            from pypdf import PdfReader
            
            reader = PdfReader(file_path)
            content_parts = []
            
            for page_num, page in enumerate(reader.pages):
                text = page.extract_text()
                if text:
                    content_parts.append(f"[Page {page_num + 1}]\n{text}")
            
            content = "\n\n".join(content_parts)
            
            # Extract metadata
            metadata = {
                "num_pages": len(reader.pages),
                "title": reader.metadata.title if reader.metadata else None,
                "author": reader.metadata.author if reader.metadata else None,
            }
            
            return Document(
                content=content,
                metadata=metadata,
                source=file_path,
                doc_type="pdf"
            )
        except ImportError:
            logger.error("pypdf not installed. Run: pip install pypdf")
            raise
        except Exception as e:
            logger.error(f"Error parsing PDF {file_path}: {e}")
            raise


class DocxParser(DocumentParser):
    """
    DOCX Document Parser
    Uses python-docx for parsing Word documents.
    """
    
    SUPPORTED_EXTENSIONS = [".docx", ".doc"]
    
    def supports(self, file_path: str) -> bool:
        return Path(file_path).suffix.lower() in self.SUPPORTED_EXTENSIONS
    
    def parse(self, file_path: str) -> Document:
        """Parse DOCX document."""
        try:
            from docx import Document as DocxDocument
            
            doc = DocxDocument(file_path)
            content_parts = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    # Check if it's a heading
                    if para.style.name.startswith("Heading"):
                        level = para.style.name.replace("Heading ", "")
                        prefix = "#" * int(level) if level.isdigit() else "##"
                        content_parts.append(f"{prefix} {para.text}")
                    else:
                        content_parts.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                table_content = []
                for row in table.rows:
                    row_content = [cell.text.strip() for cell in row.cells]
                    table_content.append(" | ".join(row_content))
                if table_content:
                    content_parts.append("\n[Table]\n" + "\n".join(table_content))
            
            content = "\n\n".join(content_parts)
            
            metadata = {
                "num_paragraphs": len(doc.paragraphs),
                "num_tables": len(doc.tables),
            }
            
            return Document(
                content=content,
                metadata=metadata,
                source=file_path,
                doc_type="docx"
            )
        except ImportError:
            logger.error("python-docx not installed. Run: pip install python-docx")
            raise
        except Exception as e:
            logger.error(f"Error parsing DOCX {file_path}: {e}")
            raise


class HTMLParser(DocumentParser):
    """
    HTML Document Parser
    Uses BeautifulSoup for parsing HTML content.
    """
    
    SUPPORTED_EXTENSIONS = [".html", ".htm"]
    
    def supports(self, file_path: str) -> bool:
        return Path(file_path).suffix.lower() in self.SUPPORTED_EXTENSIONS
    
    def parse(self, file_path: str) -> Document:
        """Parse HTML document."""
        try:
            from bs4 import BeautifulSoup
            
            with open(file_path, "r", encoding="utf-8") as f:
                html_content = f.read()
            
            soup = BeautifulSoup(html_content, "html.parser")
            
            # Remove script and style elements
            for element in soup(["script", "style", "nav", "footer", "header"]):
                element.decompose()
            
            # Extract title
            title = soup.title.string if soup.title else None
            
            # Extract main content
            main_content = soup.find("main") or soup.find("article") or soup.body or soup
            
            # Process content with structure preservation
            content_parts = []
            
            for element in main_content.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "li", "td"]):
                text = element.get_text(strip=True)
                if text:
                    if element.name.startswith("h"):
                        level = int(element.name[1])
                        content_parts.append(f"{'#' * level} {text}")
                    elif element.name == "li":
                        content_parts.append(f"• {text}")
                    else:
                        content_parts.append(text)
            
            content = "\n\n".join(content_parts)
            
            # Fallback to simple text extraction if structured extraction yields little
            if len(content) < 100:
                content = soup.get_text(separator="\n", strip=True)
            
            metadata = {
                "title": title,
                "num_links": len(soup.find_all("a")),
            }
            
            return Document(
                content=content,
                metadata=metadata,
                source=file_path,
                doc_type="html"
            )
        except ImportError:
            logger.error("beautifulsoup4 not installed. Run: pip install beautifulsoup4")
            raise
        except Exception as e:
            logger.error(f"Error parsing HTML {file_path}: {e}")
            raise


class TextParser(DocumentParser):
    """
    Plain Text Parser
    Handles .txt, .md, and other plain text files.
    """
    
    SUPPORTED_EXTENSIONS = [".txt", ".md", ".rst", ".csv", ".json", ".yaml", ".yml"]
    
    def supports(self, file_path: str) -> bool:
        return Path(file_path).suffix.lower() in self.SUPPORTED_EXTENSIONS
    
    def parse(self, file_path: str) -> Document:
        """Parse plain text document."""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
            
            # Basic metadata
            lines = content.split("\n")
            metadata = {
                "num_lines": len(lines),
                "num_chars": len(content),
                "file_extension": Path(file_path).suffix,
            }
            
            return Document(
                content=content,
                metadata=metadata,
                source=file_path,
                doc_type="text"
            )
        except Exception as e:
            logger.error(f"Error parsing text file {file_path}: {e}")
            raise


class AIParser(DocumentParser):
    """
    AI-Enhanced Document Parser
    Uses LLM to extract and structure content from documents.
    Particularly useful for complex layouts, images with text, and tables.
    """
    
    def __init__(self, llm_client=None):
        self.llm_client = llm_client
        self._base_parsers = [PDFParser(), DocxParser(), HTMLParser(), TextParser()]
    
    def supports(self, file_path: str) -> bool:
        return any(parser.supports(file_path) for parser in self._base_parsers)
    
    def parse(self, file_path: str) -> Document:
        """
        Parse document with AI enhancement.
        First extracts raw content, then uses LLM to structure and clean.
        """
        # Find appropriate base parser
        base_parser = None
        for parser in self._base_parsers:
            if parser.supports(file_path):
                base_parser = parser
                break
        
        if not base_parser:
            raise ValueError(f"No parser found for {file_path}")
        
        # Get base document
        doc = base_parser.parse(file_path)
        
        # If no LLM client, return base parsing
        if not self.llm_client:
            return doc
        
        # AI enhancement prompt
        enhancement_prompt = f"""
        You are a document processing assistant. Clean and structure the following document content:
        
        1. Fix any OCR errors or formatting issues
        2. Organize content into clear sections
        3. Preserve important information like dates, names, and numbers
        4. Remove redundant whitespace and formatting artifacts
        5. Maintain the semantic structure of the document
        
        Document content:
        {doc.content[:10000]}  # Limit to avoid token limits
        
        Return the cleaned and structured content:
        """
        
        try:
            # Call LLM for enhancement
            response = self.llm_client.generate(enhancement_prompt)
            enhanced_content = response.strip()
            
            doc.content = enhanced_content
            doc.metadata["ai_enhanced"] = True
            
        except Exception as e:
            logger.warning(f"AI enhancement failed, using base parsing: {e}")
            doc.metadata["ai_enhanced"] = False
        
        return doc


class UnstructuredParser(DocumentParser):
    """
    Parser using the Unstructured library for advanced document parsing.
    Handles complex documents with multiple content types.
    """
    
    SUPPORTED_EXTENSIONS = [".pdf", ".docx", ".doc", ".html", ".htm", ".txt", ".md", 
                           ".pptx", ".xlsx", ".eml", ".msg"]
    
    def supports(self, file_path: str) -> bool:
        return Path(file_path).suffix.lower() in self.SUPPORTED_EXTENSIONS
    
    def parse(self, file_path: str) -> Document:
        """Parse document using Unstructured library."""
        try:
            from unstructured.partition.auto import partition
            
            elements = partition(filename=file_path)
            
            content_parts = []
            element_types = set()
            
            for element in elements:
                element_types.add(type(element).__name__)
                content_parts.append(str(element))
            
            content = "\n\n".join(content_parts)
            
            metadata = {
                "num_elements": len(elements),
                "element_types": list(element_types),
            }
            
            return Document(
                content=content,
                metadata=metadata,
                source=file_path,
                doc_type=Path(file_path).suffix[1:]
            )
        except ImportError:
            logger.error("unstructured not installed. Run: pip install unstructured")
            raise
        except Exception as e:
            logger.error(f"Error parsing with Unstructured {file_path}: {e}")
            raise


def get_parser(file_path: str, use_ai: bool = False, llm_client=None) -> DocumentParser:
    """
    Factory function to get the appropriate parser for a file.
    
    Args:
        file_path: Path to the document
        use_ai: Whether to use AI-enhanced parsing
        llm_client: LLM client for AI parsing
        
    Returns:
        Appropriate DocumentParser instance
    """
    if use_ai:
        return AIParser(llm_client)
    
    parsers = [PDFParser(), DocxParser(), HTMLParser(), TextParser()]
    
    for parser in parsers:
        if parser.supports(file_path):
            return parser
    
    # Fallback to text parser
    return TextParser()


def parse_document(file_path: str, use_ai: bool = False, llm_client=None) -> Document:
    """
    Parse a single document.
    
    Args:
        file_path: Path to the document
        use_ai: Whether to use AI-enhanced parsing
        llm_client: LLM client for AI parsing
        
    Returns:
        Parsed Document object
    """
    parser = get_parser(file_path, use_ai, llm_client)
    return parser.parse(file_path)


def parse_directory(
    directory: str,
    recursive: bool = True,
    use_ai: bool = False,
    llm_client=None,
    extensions: Optional[List[str]] = None
) -> List[Document]:
    """
    Parse all documents in a directory.
    
    Args:
        directory: Path to the directory
        recursive: Whether to search subdirectories
        use_ai: Whether to use AI-enhanced parsing
        llm_client: LLM client for AI parsing
        extensions: List of file extensions to include (e.g., [".pdf", ".txt"])
        
    Returns:
        List of parsed Document objects
    """
    documents = []
    directory_path = Path(directory)
    
    if not directory_path.exists():
        raise ValueError(f"Directory not found: {directory}")
    
    # Default extensions if not specified
    if extensions is None:
        extensions = [".pdf", ".docx", ".doc", ".html", ".htm", ".txt", ".md"]
    
    # Normalize extensions
    extensions = [ext.lower() if ext.startswith(".") else f".{ext.lower()}" for ext in extensions]
    
    # Find all matching files
    pattern = "**/*" if recursive else "*"
    
    for file_path in directory_path.glob(pattern):
        if file_path.is_file() and file_path.suffix.lower() in extensions:
            try:
                doc = parse_document(str(file_path), use_ai, llm_client)
                documents.append(doc)
                logger.info(f"Parsed: {file_path}")
            except Exception as e:
                logger.error(f"Failed to parse {file_path}: {e}")
    
    logger.info(f"Parsed {len(documents)} documents from {directory}")
    return documents
